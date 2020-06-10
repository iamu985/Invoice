import os
import sys
import docx
import datetime
import pandas as pd


class Invoice:
    '''
|------------------------------------------------------------------------------------------|
|All the functions required to make this application work are defined here|
|------------------------------------------------------------------------------------------|
|=====================List of all the functions=====================|
    O-> get_files(path) = requires the path parameter i.e path of the
    directory of your docx files
        returns >>>> two lists that is document title and path of the
        documents

    O-> get_full_text(filepath) = requires the path of document
        returns >>>> parses the docx file and resturns a list of all the
        paragraphs in the docx file

    O->word_count(filepath) = same as get_full_text
        returns >>>> an integer value of the word count of that particular
        docx file

    O->all_file_wordcount(path_list) = requires a list of path of the docx
        files generally given by get_files() function
            returns >>>> a list of integer values of word count in each docx
            file

    O-> dummy_wordcount(word_counts_list) = requires a list of integer
        wordcounts, generally given by all_file_wordcount
            returns >>>> a list of rounded off values or assumption
                wordcount for a docx file

    O-> getProperties(path_list)

    '''
    def get_files(path):
        '''
        returns a list of docx file name and path
        '''
        
        docx_file, file_path = [], []
        
        for r, d, f in os.walk(path):
            for file in f:
                if '.docx' in file:
                    docx_file.append(file)
                    file_path.append(os.path.join(r, file))
        
        return docx_file, file_path

    def get_full_text(filepath):
        '''
        function to get the list of paragraph from the docx file
        '''
        full_text = []
        document = docx.Document(filepath)
        content_paragraph = document.paragraphs
        for paragraph in content_paragraph:
            full_text.append(paragraph.text)
        
        return full_text

    #functions required to get the values that makes up a dataframe
    def word_count (filepath):
        '''
        returns the word count of a particular file
        '''
        context = Invoice.get_full_text(filepath)
        i = 0
        paragraph_count = []
        while i < len(context):
            text_data = context[i].split(' ')
            if len(text_data) != 1:
                if len(text_data) != 2:
                    paragraph_count.append(len(text_data))
            i+=1
        
        return sum(paragraph_count) - 10

    def all_file_wordcount(path_list):
        '''
        returns a list of word count for all files
        '''
        word_counts = []
        for i in range(len(path_list)):
            counts = Invoice.word_count(path_list[i])
            word_counts.append(counts)
        
        return word_counts

    def dummy_wordcount(wordcounts_list):
        '''
        returns a list of round off figure of target word counts
        '''
        dummy_count = []
        for count in wordcounts_list:
            if count > 1000:
                diff = count%1000
            if count < 1000:
                diff = count%100
            
            value = count - diff
            dummy_count.append(value)
        
        return dummy_count

    def getProperties(path_list):
        '''
        returns a list of date and string values of file created, modified and file title
        '''
        created, modified, title = [], [], []
        for i in range(len(path_list)):
            docs = docx.Document(path_list[i])
            properties = docs.core_properties
            
            #title of the file
            title_file = properties.title
            title.append(title_file)
            
            #file created date
            date_created = properties.created
            created_date_format = str(date_created.year) + '/' + str(date_created.month) + '/' + str(date_created.day)
            created.append(created_date_format)
            
            #file modified date
            date_modified = properties.modified
            modified_date_format = str(date_modified.year) + '/' + str(date_modified.month) + '/' + str(date_modified.day)
            modified.append(modified_date_format)
            
        return created, modified, title

    def proper_properties(filename, title, created, modified, wordcount, dummycount, title_bool=False, dummy_bool=False):
        '''returns a list of properties'''
        
        if title_bool==False and dummy_bool==False:
            proper_values = [filename, created, modified, wordcount]
            return proper_values
        if title_bool==True and dummy_bool==False:
            proper_values = [filename, title, created, modified, wordcount]
            return proper_values
        if title_bool==False and dummy_bool==True:
            proper_values = [filename, created, modified, wordcount, dummycount]
            return proper_values
        if title_bool==True and dummy_bool==True:
            proper_values = [filename, title, created, modified, wordcount, dummycount]
            return proper_values

        
            

    def dictionary(values, title_bool=False, dummy_bool=False):
        '''
        returns dictionary data
        '''
        data = {}
        if title_bool == False and dummy_bool == False:
            keys = ['file_name', 'created', 'modified', 'word_count']
            for k, v in zip(keys, values):
                data[k] = v
            return data
            
        if title_bool == True and dummy_bool == False:
            keys = ['file_name', 'title', 'created', 'modified', 'word_count']
            for k, v in zip(keys, values):
                data[k] = v
            return data
        
        if dummy_bool == True and title_bool == False:
            keys = ['file_name', 'created', 'modified', 'word_count', 'dummy_count']
            for k, v in zip(keys, values):
                data[k] = v
            return data
        
        if dummy_bool == True and title_bool == True:
            keys = ['file_name', 'title', 'created', 'modified', 'word_count', 'dummy_count']
            for k, v in zip(keys, values):
                data[k] = v
            return data


    def table(data):
        '''
        function to draw the table of data
        '''
        table = pd.DataFrame(data)
        return table

    def export(table, to_path, save_file_name):
        '''
        function to export the data into an excel file
        '''
        sep = '/'
        ext = '.xlsx'
        #excel writer
        path = to_path + sep + save_file_name + ext
        with pd.ExcelWriter(path, date_format='YYYY/MM/DD') as writer:
            table.to_excel(writer, index=False, header=True)
        
        



